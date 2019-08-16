from docx import Document

def create_document(report_data):
    name = ""
    if 'Individual Identifiers' in report_data and 'Name' in report_data['Individual Identifiers']:
        name = report_data['Individual Identifiers']['Name']
    doc = Document()
    doc.add_heading("Individual Identifiers", level=9)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Identifier'
    hdr_cells[1].text = 'Verified Information'
    hdr_cells[2].text = 'Source Reliability'
    row_cells = table.add_row().cells
    row_cells[0].text = "Given name"
    row_cells[1].text = name
    row_cells = table.add_row().cells
    row_cells[0].text = "Alias"
    row_cells = table.add_row().cells
    row_cells[0].text = "Gender"
    row_cells = table.add_row().cells
    row_cells[0].text = "Date of birth"
    row_cells[1].text = report_data['Individual Identifiers']['Date of Birth']
    row_cells = table.add_row().cells
    row_cells[0].text = "Country of domicile"
    row_cells = table.add_row().cells
    row_cells[0].text = "Nationality"
    row_cells = table.add_row().cells
    row_cells[0].text = "Current address"
    row_cells[1].text = report_data['Individual Identifiers']['Address']
    doc.add_paragraph('')
    doc.add_heading("Corporate Interests", level=9)
    doc.add_heading("Active Company Appointments", level=9)
    table = add_company_table(doc, report_data['Current Directorships'])
    doc.add_paragraph('')
    doc.add_heading("Previous Company Appointments", level=9)
    table = add_company_table(doc, report_data['Past Directorships'])
    doc.add_paragraph('')
    doc.add_heading("Known Associates", level=9)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Relationship'
    hdr_cells[2].text = 'Source Reliability'

    row_cells = table.add_row().cells
    row_cells[0].text = "Family"

    for fam in report_data['Family']:
        row_cells = table.add_row().cells
        row_cells[0].text = fam

    row_cells = table.add_row().cells
    row_cells[0].text = "Business"

    for associate in report_data['Codirectors']:
        row_cells = table.add_row().cells
        row_cells[0].text = associate
    doc.add_paragraph('')
    doc.add_heading("Risk Profiling", level=9)
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Target'
    hdr_cells[1].text = 'Sanctions'
    hdr_cells[2].text = 'Political Exposure'
    hdr_cells[3].text = """Regulatory Enforcement/Legal/Credit Issues"""
    hdr_cells[4].text = """Significant Adverse News"""
    hdr_cells[5].text = """High-Risk Jurisdiction"""
    hdr_cells[6].text = """High-Risk Industry"""
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells = table.add_row().cells
    row_cells[0].text = "Connected Entities"
    for target in report_data['Risk Profiling']:
        row_cells = table.add_row().cells
        row_cells[0].text = target
    filename = "{}.docx".format(name)
    doc.save(filename)
    return filename

def add_company_table(doc, company_list):
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Company Name'
    hdr_cells[1].text = 'Country of Registration'
    hdr_cells[2].text = 'Position'
    hdr_cells[3].text = 'Source Reliability'
    for company in company_list:
        row_cells = table.add_row().cells
        row_cells[0].text = company
        row_cells[2].text = "Director"

    return table

def parse_company_pdf(pdf):
    for page in pdf:
        lines = page.split("\n")
        for line in lines:
            print(line)
    return {'Name': 'Company Name'}

def parse_person_pdf(pdf):
    data = {}
    all_lines = [] 
    for page in pdf:
        lines = page.split("\n")
        filtered_lines = [ x for x in lines if x[0:4] != "https" ]
        all_lines += filtered_lines

    line_iter = iter(all_lines)
    for line in line_iter:
        words = line.split()
        if len(words) == 1 and words[0] == "Residency":
            break

    ident = {}
    #Parse individual identifiers
    for line in line_iter:
        title = line.strip()
        if title == "Company Directorships":
            break
        
        kvp = line.split(":")
        if len(kvp) > 1:
            ident[kvp[0].strip()] = kvp[1].strip()

    data['Individual Identifiers'] = ident
    current_companies = []
    past_companies = []
    current_directorship = False
    last_company = ""
    #Parse Company Directorships 
    for line in line_iter:
        title = line.strip()
        if title == "Legal":
            break

        kvp = line.split(":")
        if len(kvp) > 1:
            key = kvp[0].strip()
            value = kvp[1].strip()
            if key == "Date Resigned" and value == "Unavailable":
                current_directorship = True
            if key == "Company Name":
                last_company = value
            if key == "Last Filed Annual Return":
                if value != "Unavailable":                #make sure company is still active
                    year = int(value[-4:])
                    if year < 2018:
                        current_directorship = False
                if last_company != "":
                    if current_directorship:
                        current_companies.append(last_company)
                    else:
                        past_companies.append(last_company)
                current_directorship = False
        
    data['Current Directorships'] = current_companies
    data['Past Directorships'] = past_companies

    for line in line_iter:
        words = line.split()
        if len(words) == 6 and words[2] == "Phonematch":
            break

    family_members = []
    last_name = ident['Name'].split()[-1]
    #Parse Family Members from Occupants 
    for line in line_iter:
        words = line.split()
        if len(words) == 2 and words[0] == "Companies" and words[1] == "House":
            break
        name = ""
        for word in words:
            if word[0].isdigit():
                break
            name += "{} ".format(word)
            if word == last_name:
                family_members.append(name.strip())
                break
            
    data['Family'] = family_members

    for line in line_iter:
        words = line.split()
        if len(words) > 2 and words[-1] == "directorship":
            break

    #Parse Co-directors
    directors = []
    for line in line_iter:
        title = line.strip()
        if title == "Personal Associates":
            break
        words = line.split()
        if len(words) > 3 and words[-1] == "Current":
            if words[-3][-1] == ",":
                first_name = words[-4][0:-1]
                second_name = words[-3][0:-1]
                last_name = next(line_iter).strip()
            else:
                last_name = words[-3]
                second_name = words[-4][0:-1]
                first_name = ""
                if len(words) > 4 and words[-5][-1] == ",":
                    first_name = words[-5][0:-1]
            
            director_name = "{} {} {}".format(first_name, second_name, last_name).strip()
            if director_name not in directors:
                directors.append(director_name)

    data['Codirectors'] = directors
    
    data['Risk Profiling'] = family_members + directors + current_companies

    return data
